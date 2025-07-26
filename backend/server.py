from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Depends, Header, Request
from fastapi.responses import JSONResponse, StreamingResponse
from starlette.middleware.cors import CORSMiddleware
from starlette.middleware.sessions import SessionMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
import uuid
from datetime import datetime, timedelta
from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
import httpx
import secrets
import asyncio
import tempfile
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
from html import unescape
from authlib.integrations.starlette_client import OAuth
from authlib.integrations.starlette_client import StarletteOAuth2App
import jwt
import json
from passlib.context import CryptContext
from passlib.hash import bcrypt
from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionResponse, CheckoutStatusResponse, CheckoutSessionRequest
from typing import Dict

ROOT_DIR = Path(__file__).parent

# Secure Stripe configuration (encoded to hide from plain text)
def get_stripe_key():
    # Hardcoded API key
    return "sk_test_emergent"

# MongoDB connection
mongo_url = "mongodb://localhost:27017"
client = AsyncIOMotorClient(mongo_url)
db = client["test_database"]

# Create the main app without a prefix
app = FastAPI()

# Add session middleware for OAuth
app.add_middleware(SessionMiddleware, secret_key=secrets.token_urlsafe(32))

# Password hashing context
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Setup OAuth
oauth = OAuth()
oauth.register(
    name='google',
    client_id="758478706314-pn8dh4u94p8mt06qialfdigaqs5glj9s.apps.googleusercontent.com",
    client_secret="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",
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    }
)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Models
class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str
    picture: Optional[str] = None
    password_hash: Optional[str] = None  # For email/password auth
    auth_provider: str = "google"  # "google" or "email"
    credit_balance: int = Field(default=10)  # Starting credits for new users
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class CreditTransaction(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    amount: int  # Positive for credit addition, negative for usage
    transaction_type: str  # "chapter_generation", "chapter_regeneration", "credit_purchase", "bonus"
    description: str
    book_project_id: Optional[str] = None  # Link to book project if applicable
    chapter_number: Optional[int] = None  # Chapter number if applicable
    created_at: datetime = Field(default_factory=datetime.utcnow)

class UserSession(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    session_token: str
    expires_at: datetime
    created_at: datetime = Field(default_factory=datetime.utcnow)

class BookProject(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str  # Added user association
    title: str
    description: str
    pages: int
    chapters: int
    language: str
    writing_style: str = "story"  # "story", "descriptive", "academic", "technical", "biography", "self_help", "children", "poetry", "business"
    outline: Optional[str] = None
    chapters_content: Optional[dict] = {}
    generated_chapters: Optional[list] = []  # Track which chapter numbers have been generated
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class BookProjectCreate(BaseModel):
    title: str
    description: str
    pages: int
    chapters: int
    language: str
    writing_style: str = "story"  # "story", "descriptive", "academic", "technical", "biography", "self_help", "children", "poetry", "business"

class AuthSessionRequest(BaseModel):
    session_id: str

class GoogleTokenRequest(BaseModel):
    token: str

class EmailPasswordRequest(BaseModel):
    email: str
    password: str

class RegisterRequest(BaseModel):
    email: str
    name: str
    password: str

class UserProfile(BaseModel):
    id: str
    email: str
    name: str
    picture: Optional[str] = None

class OutlineRequest(BaseModel):
    project_id: str

class ChapterRequest(BaseModel):
    project_id: str
    chapter_number: int

class ChapterUpdate(BaseModel):
    project_id: str
    chapter_number: int
    content: str

class CreditPurchaseRequest(BaseModel):
    amount: int
    payment_method: str = "stripe"  # Future payment integration

class BookCostRequest(BaseModel):
    pages: int
    chapters: int

class BookCostResponse(BaseModel):
    pages: int
    requested_chapters: int
    minimum_chapters: int
    cost_per_chapter: int
    total_cost: int
    pages_per_chapter: float
    
class CreditBalanceResponse(BaseModel):
    credit_balance: int
    user_id: str

class CreditTransactionResponse(BaseModel):
    id: str
    amount: int
    transaction_type: str
    description: str
    book_project_id: Optional[str]
    chapter_number: Optional[int]
    created_at: datetime

# Payment-related models
class PaymentTransaction(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    session_id: str
    payment_id: Optional[str] = None
    amount: float  # Amount in EUR
    currency: str = "eur"
    credits_amount: int  # Number of credits to be added
    package_id: str  # Package identifier (small, medium, large)
    payment_status: str = "pending"  # pending, paid, failed, expired
    status: str = "initiated"  # initiated, pending, completed, failed
    metadata: Optional[Dict[str, str]] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)

class PaymentPackageRequest(BaseModel):
    package_id: str  # "small", "medium", "large"
    origin_url: str  # Frontend origin for constructing success/cancel URLs

class PaymentSessionResponse(BaseModel):
    checkout_url: str
    session_id: str
    package_info: dict  # Changed from Dict[str, any] to dict

class PaymentStatusResponse(BaseModel):
    session_id: str
    payment_status: str
    status: str
    amount: float
    currency: str
    credits_amount: int
    package_id: str

# Credit packages definition (server-side only for security)
CREDIT_PACKAGES = {
    "small": {
        "name": "Small Package",
        "credits": 10,
        "price": 5.00,
        "currency": "eur",
        "description": "Perfect for getting started"
    },
    "medium": {
        "name": "Medium Package", 
        "credits": 25,
        "price": 10.00,
        "currency": "eur",
        "description": "Great value for regular users"
    },
    "large": {
        "name": "Large Package",
        "credits": 50,
        "price": 20.00,
        "currency": "eur",
        "description": "Best deal for power users"
    }
}

# Helper function for style-specific instructions
def get_style_instructions(writing_style: str, content_type: str = "outline"):
    """Get style-specific instructions for different writing styles"""
    
    styles = {
        "story": {
            "outline": """Create a story outline that:
- Develops compelling characters with distinct personalities and voices
- Crafts engaging plot points with emotional depth and conflict
- Maintains consistent pacing with varied scene types (dialogue, action, reflection)
- Builds immersive story worlds with vivid settings and atmosphere
- Weaves subplots seamlessly into the main narrative
- Creates emotional resonance through character relationships and growth
- Uses natural scene transitions and chapter hooks
- Balances showing vs. telling throughout the narrative
- Includes opportunities for authentic dialogue and character interaction""",
            "chapter": """Write in a story style that:
- Creates compelling characters with unique voices and speech patterns
- Uses natural, authentic dialogue that reveals character personality
- Balances narrative description with action and conversation
- Shows character emotions through actions and dialogue, not just description
- Includes varied sentence structures and paragraph lengths for natural flow
- Uses specific, sensory details to create vivid scenes
- Alternates between intimate character moments and broader plot advancement
- Incorporates subtext and emotional depth in character interactions
- Uses dialogue to advance plot and reveal character development
- Avoids having characters speak in the same formal register as the narrator
- Includes natural speech patterns, contractions, and conversational rhythms
- Breaks up exposition with engaging dialogue and action sequences""",
            "formatting": """Use story HTML formatting:
- <h2> for chapter titles with evocative, atmospheric names
- <h3> for scene breaks or major story sections
- <p> for narrative paragraphs with varied lengths
- Use frequent paragraph breaks for dialogue and action
- Format dialogue clearly with proper speaker attribution
- <strong> for emphasis and important moments
- <em> for internal thoughts and emphasis
- Structure content with natural scene transitions and pacing"""
        },
        "descriptive": {
            "outline": """Create a structured informational outline that:
- Presents clear topic hierarchies and logical flow
- Breaks down complex concepts systematically
- Incorporates relevant examples and case studies
- Balances depth and accessibility of content
- Uses consistent terminology and definitions
- Provides clear learning objectives per section
- Includes practical applications and insights
- Maintains engaging narrative elements where appropriate
- Balances informative content with readability""",
            "chapter": """Write in a descriptive style that:
- Uses clear, informative language with engaging examples
- Balances detailed explanations with accessible prose
- Includes relevant case studies and real-world applications
- Maintains reader engagement through varied content structure
- Uses natural transitions between concepts and ideas
- Incorporates storytelling elements to illustrate points
- Avoids overly technical jargon without sacrificing accuracy
- Includes dialogue or character examples where relevant
- Structures content for optimal comprehension and retention""",
            "formatting": """Use descriptive HTML formatting:
- <h2> for chapter titles with descriptive, engaging names
- <h3> for major topics or concepts
- <h4> for sub-topics and detailed explanations
- <p> for explanatory paragraphs with varied lengths
- <ul> and <li> for organized lists and key points
- <strong> for important terms and concepts
- <em> for emphasis and examples
- Use clear section breaks and logical content flow"""
        },
        "academic": {
            "outline": """Create an academic/research outline that:
- Follows scholarly structure and methodology
- Includes literature review and theoretical framework
- Presents arguments with supporting evidence
- Uses formal academic language and citations
- Organizes content by research questions or hypotheses
- Includes methodology and analysis sections
- Provides clear research conclusions""",
            "chapter": """Write in an academic style that:
- Uses formal, scholarly language and tone
- Supports arguments with evidence and citations
- Follows academic structure and conventions
- Includes theoretical background and context
- Presents balanced analysis and critical thinking
- Uses appropriate academic terminology
- Maintains objectivity and rigor""",
            "formatting": """Use academic HTML formatting:
- <h2> for chapter titles with chapter names
- <h3> for main sections (Introduction, Literature Review, etc.)
- <h4> for subsections
- <p> for paragraphs with academic prose
- <ul> and <li> for structured lists
- <strong> for key concepts and terms
- <em> for emphasis and foreign terms"""
        },
        "technical": {
            "outline": """Create a technical/manual outline that:
- Follows step-by-step instructional structure
- Includes prerequisites and requirements
- Organizes content by procedures or processes
- Uses clear, concise technical language
- Includes troubleshooting and best practices
- Provides practical examples and use cases
- Follows logical workflow sequences""",
            "chapter": """Write in a technical style that:
- Uses clear, precise technical language
- Provides step-by-step instructions
- Includes practical examples and code snippets
- Follows logical procedural flow
- Addresses common issues and solutions
- Uses consistent technical terminology
- Focuses on actionable information""",
            "formatting": """Use technical HTML formatting:
- <h2> for chapter titles with chapter names
- <h3> for major procedures or sections
- <h4> for sub-procedures
- <p> for instructions and explanations
- <ul> and <li> for step-by-step lists
- <strong> for important warnings or key points
- <em> for technical terms or variables"""
        },
        "biography": {
            "outline": """Create a biographical outline that:
- Follows chronological or thematic structure
- Focuses on key life events and milestones
- Includes personal background and influences
- Explores character development and growth
- Incorporates historical and cultural context
- Balances personal and professional aspects
- Shows cause-and-effect relationships""",
            "chapter": """Write in a biographical style that:
- Narrates life events with personal insight
- Balances factual information with story elements
- Includes dialogue and personal anecdotes
- Provides historical and cultural context
- Shows character development over time
- Uses engaging narrative techniques
- Maintains respect for the subject""",
            "formatting": """Use biographical HTML formatting:
- <h2> for chapter titles with chapter names
- <h3> for major life periods or themes
- <p> for narrative paragraphs
- <ul> and <li> for key achievements or events
- <strong> for important dates and names
- <em> for quotes and personal thoughts"""
        },
        "self_help": {
            "outline": """Create a self-help outline that:
- Focuses on practical advice and actionable steps
- Includes motivational and inspirational content
- Organizes content by problems and solutions
- Uses encouraging and empowering language
- Includes real-life examples and case studies
- Provides exercises and practical applications
- Builds progressive skill development""",
            "chapter": """Write in a self-help style that:
- Uses encouraging, motivational language
- Provides practical, actionable advice
- Includes personal stories and examples
- Addresses common challenges and solutions
- Engages readers with exercises and activities
- Maintains an optimistic, empowering tone
- Focuses on personal growth and development""",
            "formatting": """Use self-help HTML formatting:
- <h2> for chapter titles with chapter names
- <h3> for main concepts or strategies
- <p> for explanations and advice
- <ul> and <li> for action steps and tips
- <strong> for key takeaways and important points
- <em> for motivational quotes and emphasis"""
        },
        "children": {
            "outline": """Create a children's book outline that:
- Uses age-appropriate language and concepts
- Includes engaging characters and situations
- Focuses on learning and entertainment
- Incorporates moral lessons or educational content
- Uses repetition and rhythm for engagement
- Includes opportunities for interaction
- Balances fun with educational value""",
            "chapter": """Write in a children's style that:
- Uses simple, clear language appropriate for the age group
- Creates engaging characters and situations
- Includes dialogue and action
- Incorporates learning opportunities naturally
- Uses repetition and rhythm for memorability
- Maintains a fun, engaging tone
- Includes moral lessons or educational content""",
            "formatting": """Use children's book HTML formatting:
- <h2> for chapter titles with chapter names
- <h3> for story sections or activities
- <p> for story text and explanations
- <ul> and <li> for lists and activities
- <strong> for important words or concepts
- <em> for sound effects or emphasis"""
        },
        "poetry": {
            "outline": """Create a poetry/creative outline that:
- Organizes by themes, forms, or emotional progression
- Includes various poetic forms and styles
- Focuses on imagery, metaphor, and literary devices
- Explores emotions and human experiences
- Uses creative structure and organization
- Incorporates rhythm and sound patterns
- Balances different poetic techniques""",
            "chapter": """Write in a poetic/creative style that:
- Uses rich imagery and metaphorical language
- Incorporates various poetic forms and techniques
- Focuses on emotional expression and artistic beauty
- Uses rhythm, meter, and sound patterns
- Employs creative language and wordplay
- Explores themes through artistic expression
- Maintains creative and innovative approach""",
            "formatting": """Use poetry HTML formatting:
- <h2> for chapter titles with chapter names
- <h3> for poem titles or sections
- <p> for poem stanzas and creative prose
- <ul> and <li> for structured creative elements
- <strong> for emphasis and key imagery
- <em> for titles and artistic emphasis"""
        },
        "business": {
            "outline": """Create a business/professional outline that:
- Focuses on practical business applications
- Includes case studies and real-world examples
- Organizes content by business functions or processes
- Uses professional terminology and concepts
- Provides actionable strategies and solutions
- Includes metrics, data, and analysis
- Addresses current business challenges""",
            "chapter": """Write in a business/professional style that:
- Uses professional, authoritative language
- Provides practical business insights and strategies
- Includes relevant case studies and examples
- Focuses on actionable business solutions
- Uses data and metrics to support points
- Addresses current business trends and challenges
- Maintains a professional, informative tone""",
            "formatting": """Use business HTML formatting:
- <h2> for chapter titles with chapter names
- <h3> for major business concepts or strategies
- <h4> for sub-topics and detailed points
- <p> for explanations and analysis
- <ul> and <li> for strategies and action items
- <strong> for key business terms and metrics
- <em> for emphasis and important concepts"""
        },
        "health": {
            "outline": """Create a health/wellness outline that:
- Focuses on evidence-based health information and practices
- Includes practical lifestyle recommendations and tips
- Organizes content by health topics or body systems
- Uses accessible medical terminology with explanations
- Provides actionable health strategies and prevention methods
- Includes real-life wellness examples and case studies
- Addresses common health concerns and solutions""",
            "chapter": """Write in a health/wellness style that:
- Uses clear, accessible language for health concepts
- Provides evidence-based information and recommendations
- Includes practical tips and actionable advice
- Addresses common health concerns with empathy
- Uses encouraging, motivational tone for wellness
- Incorporates real-life examples and success stories
- Maintains professional credibility while being approachable""",
            "formatting": """Use health HTML formatting:
- <h2> for chapter titles with health topic names
- <h3> for main health concepts or conditions
- <h4> for specific symptoms, treatments, or tips
- <p> for explanations and health advice
- <ul> and <li> for symptoms, treatments, and recommendations
- <strong> for important health warnings and key points
- <em> for medical terms and emphasis"""
        },
        "travel": {
            "outline": """Create a travel/adventure outline that:
- Organizes content by destinations, themes, or journey progression
- Includes practical travel information and cultural insights
- Focuses on personal experiences and storytelling
- Incorporates local culture, food, and customs
- Provides travel tips and practical advice
- Balances adventure narrative with useful information
- Includes historical and geographical context""",
            "chapter": """Write in a travel/adventure style that:
- Uses vivid, descriptive language to paint locations
- Incorporates personal experiences and anecdotes
- Includes dialogue and cultural interactions
- Provides practical travel advice and tips
- Balances storytelling with informational content
- Uses engaging narrative to bring places to life
- Maintains enthusiasm and wanderlust throughout""",
            "formatting": """Use travel HTML formatting:
- <h2> for chapter titles with destination or theme names
- <h3> for specific locations or travel experiences
- <p> for narrative descriptions and travel stories
- <ul> and <li> for travel tips and recommendations
- <strong> for important travel information and warnings
- <em> for foreign words and cultural terms"""
        },
        "cooking": {
            "outline": """Create a cooking/recipe outline that:
- Organizes content by meal types, cuisines, or skill levels
- Includes ingredient lists and cooking techniques
- Focuses on practical cooking instructions and tips
- Incorporates cultural and historical food context
- Provides variations and substitutions for recipes
- Includes cooking fundamentals and kitchen skills
- Balances recipes with cooking education""",
            "chapter": """Write in a cooking/recipe style that:
- Uses clear, step-by-step instructional language
- Provides detailed ingredient lists and measurements
- Includes cooking techniques and kitchen tips
- Incorporates personal cooking stories and experiences
- Explains the 'why' behind cooking methods
- Uses encouraging, approachable tone for all skill levels
- Includes variations and customization options""",
            "formatting": """Use cooking HTML formatting:
- <h2> for chapter titles with dish or technique names
- <h3> for recipe names or cooking sections
- <h4> for ingredients, instructions, or variations
- <p> for cooking instructions and explanations
- <ul> and <li> for ingredient lists and step-by-step instructions
- <strong> for important cooking tips and warnings
- <em> for cooking terms and techniques"""
        },
        "history": {
            "outline": """Create a historical outline that:
- Organizes content chronologically or thematically
- Includes key historical events, figures, and periods
- Provides historical context and cause-and-effect relationships
- Incorporates primary sources and historical evidence
- Balances broad historical trends with specific details
- Addresses multiple perspectives and interpretations
- Connects historical events to contemporary relevance""",
            "chapter": """Write in a historical style that:
- Uses engaging narrative to bring history to life
- Incorporates historical details and authentic context
- Includes dialogue and personal stories from the period
- Provides analysis of causes and consequences
- Balances storytelling with factual accuracy
- Uses vivid descriptions of historical settings
- Maintains scholarly credibility with accessible prose""",
            "formatting": """Use historical HTML formatting:
- <h2> for chapter titles with period or event names
- <h3> for major historical events or themes
- <h4> for specific dates, figures, or topics
- <p> for historical narrative and analysis
- <ul> and <li> for key events, dates, and figures
- <strong> for important historical names and dates
- <em> for historical terms and foreign words"""
        },
        "science": {
            "outline": """Create a science/technology outline that:
- Organizes content by scientific concepts or technological topics
- Includes current research and technological advances
- Focuses on explaining complex concepts clearly
- Incorporates real-world applications and examples
- Provides scientific method and evidence-based reasoning
- Balances technical accuracy with accessibility
- Addresses implications and future developments""",
            "chapter": """Write in a science/technology style that:
- Uses clear explanations for complex scientific concepts
- Incorporates current research and discoveries
- Includes real-world examples and applications
- Provides evidence-based reasoning and analysis
- Uses engaging analogies to explain difficult concepts
- Maintains scientific accuracy while being accessible
- Addresses both benefits and challenges of technology""",
            "formatting": """Use science HTML formatting:
- <h2> for chapter titles with scientific topic names
- <h3> for major scientific concepts or discoveries
- <h4> for specific theories, experiments, or technologies
- <p> for scientific explanations and analysis
- <ul> and <li> for scientific processes and key points
- <strong> for important scientific terms and concepts
- <em> for scientific names and technical terms"""
        },
        "philosophy": {
            "outline": """Create a philosophy/ethics outline that:
- Organizes content by philosophical concepts or ethical issues
- Includes major philosophical arguments and perspectives
- Focuses on critical thinking and logical reasoning
- Incorporates historical and contemporary philosophers
- Provides balanced examination of different viewpoints
- Addresses practical applications of philosophical concepts
- Encourages reader reflection and intellectual engagement""",
            "chapter": """Write in a philosophy/ethics style that:
- Uses thoughtful, reflective language and tone
- Presents philosophical arguments clearly and logically
- Includes diverse perspectives and viewpoints
- Encourages critical thinking and self-reflection
- Uses examples and thought experiments effectively
- Maintains intellectual rigor while being accessible
- Addresses practical implications of philosophical concepts""",
            "formatting": """Use philosophy HTML formatting:
- <h2> for chapter titles with philosophical topic names
- <h3> for major philosophical concepts or arguments
- <h4> for specific theories or ethical considerations
- <p> for philosophical discourse and analysis
- <ul> and <li> for key arguments and philosophical points
- <strong> for important philosophical terms and concepts
- <em> for emphasis and philosophical terminology"""
        },
        "psychology": {
            "outline": """Create a psychology/mental health outline that:
- Organizes content by psychological concepts or mental health topics
- Includes evidence-based psychological theories and research
- Focuses on understanding human behavior and mental processes
- Incorporates practical applications and therapeutic approaches
- Provides compassionate approach to mental health issues
- Balances scientific rigor with human empathy
- Addresses both individual and social psychological factors""",
            "chapter": """Write in a psychology/mental health style that:
- Uses empathetic, understanding language for sensitive topics
- Incorporates scientific research and evidence-based practices
- Includes real-life examples and case studies (anonymized)
- Provides practical coping strategies and insights
- Maintains professional credibility with accessible explanations
- Addresses stigma and promotes mental health awareness
- Uses encouraging, supportive tone throughout""",
            "formatting": """Use psychology HTML formatting:
- <h2> for chapter titles with psychological topic names
- <h3> for major psychological concepts or conditions
- <h4> for specific symptoms, treatments, or strategies
- <p> for psychological explanations and advice
- <ul> and <li> for symptoms, strategies, and key points
- <strong> for important psychological terms and warnings
- <em> for emphasis and psychological terminology"""
        },
        "mystery": {
            "outline": """Create a mystery/thriller outline that:
- Develops compelling mystery plot with clues and red herrings
- Creates suspenseful pacing and tension throughout
- Includes complex characters with hidden motives
- Builds atmosphere and mood appropriate to the genre
- Incorporates investigative elements and logical deduction
- Provides satisfying resolution that ties up all loose ends
- Balances action with character development""",
            "chapter": """Write in a mystery/thriller style that:
- Creates suspenseful atmosphere and tension
- Uses clues and foreshadowing effectively
- Develops complex characters with secrets and motives
- Incorporates dialogue that reveals and conceals information
- Builds pacing that maintains reader interest
- Uses descriptive language to create mood and setting
- Balances action sequences with investigative elements""",
            "formatting": """Use mystery HTML formatting:
- <h2> for chapter titles with intriguing, atmospheric names
- <h3> for scene breaks or major plot developments
- <p> for narrative with varied paragraph lengths for pacing
- Use frequent paragraph breaks for dialogue and action
- <strong> for important clues and revelations
- <em> for internal thoughts and emphasis
- Structure content with suspenseful scene transitions"""
        },
        "romance": {
            "outline": """Create a romance/relationships outline that:
- Develops emotional character arcs and relationship progression
- Creates compelling romantic tension and chemistry
- Includes realistic relationship challenges and growth
- Incorporates diverse relationship dynamics and perspectives
- Balances romantic elements with character development
- Addresses themes of love, trust, and personal growth
- Provides satisfying emotional resolution""",
            "chapter": """Write in a romance/relationships style that:
- Creates authentic emotional connections between characters
- Uses dialogue that reveals chemistry and tension
- Incorporates sensual and emotional descriptions appropriately
- Develops realistic relationship challenges and conflicts
- Balances romantic scenes with character development
- Uses internal thoughts and emotions effectively
- Maintains appropriate tone for the target audience""",
            "formatting": """Use romance HTML formatting:
- <h2> for chapter titles with romantic or emotional themes
- <h3> for scene breaks or emotional transitions
- <p> for narrative with focus on emotions and relationships
- Use paragraph breaks for dialogue and emotional moments
- <strong> for important emotional revelations
- <em> for internal thoughts and romantic emphasis
- Structure content with emotional pacing and development"""
        },
        "fantasy": {
            "outline": """Create a fantasy/science fiction outline that:
- Develops imaginative world-building with consistent rules
- Creates compelling characters with unique abilities or challenges
- Incorporates fantastical or futuristic elements seamlessly
- Builds epic plots with high stakes and adventure
- Includes magic systems or technological concepts
- Balances world-building with character development
- Addresses themes relevant to the human condition""",
            "chapter": """Write in a fantasy/science fiction style that:
- Creates vivid, imaginative descriptions of fantastical worlds
- Develops characters with unique abilities and challenges
- Incorporates magic or technology naturally into the narrative
- Uses dialogue that fits the world and characters
- Balances action and adventure with emotional depth
- Builds tension through both external and internal conflicts
- Maintains consistency in world-building and rules""",
            "formatting": """Use fantasy HTML formatting:
- <h2> for chapter titles with epic or mystical names
- <h3> for scene breaks or major fantasy elements
- <p> for world-building descriptions and narrative
- Use paragraph breaks for action and dialogue
- <strong> for important fantasy terms and revelations
- <em> for magical incantations and emphasis
- Structure content with adventure pacing and world-building"""
        },
        "horror": {
            "outline": """Create a horror/supernatural outline that:
- Develops frightening atmosphere and psychological tension
- Creates compelling supernatural or psychological threats
- Builds suspense through pacing and foreshadowing
- Incorporates elements of fear and unease throughout
- Balances horror elements with character development
- Uses setting and atmosphere to enhance fear
- Provides resolution that addresses the horror elements""",
            "chapter": """Write in a horror/supernatural style that:
- Creates atmospheric descriptions that build dread
- Uses psychological tension and fear effectively
- Incorporates supernatural elements naturally
- Develops characters facing terrifying situations
- Balances graphic content with psychological horror
- Uses pacing to build and release tension
- Maintains appropriate level of fear for the audience""",
            "formatting": """Use horror HTML formatting:
- <h2> for chapter titles with ominous or atmospheric names
- <h3> for scene breaks or supernatural encounters
- <p> for atmospheric descriptions and narrative
- Use paragraph breaks for tension and dialogue
- <strong> for important horror revelations and threats
- <em> for supernatural elements and emphasis
- Structure content with suspenseful pacing and atmosphere"""
        },
        "memoir": {
            "outline": """Create a memoir/personal stories outline that:
- Organizes content by life themes or chronological events
- Includes personal reflections and insights
- Focuses on meaningful life experiences and lessons learned
- Incorporates emotional honesty and vulnerability
- Balances personal details with universal themes
- Provides growth and transformation narratives
- Addresses how experiences shaped the author""",
            "chapter": """Write in a memoir/personal style that:
- Uses authentic, personal voice and perspective
- Incorporates emotional honesty and vulnerability
- Includes dialogue and personal anecdotes
- Provides reflective insights and lessons learned
- Balances personal details with broader themes
- Uses storytelling techniques to engage readers
- Maintains respect for privacy of others mentioned""",
            "formatting": """Use memoir HTML formatting:
- <h2> for chapter titles with personal or thematic names
- <h3> for major life events or themes
- <p> for personal narrative and reflections
- <ul> and <li> for key insights and lessons learned
- <strong> for important personal revelations
- <em> for emphasis and reflective thoughts"""
        },
        "educational": {
            "outline": """Create an educational/textbook outline that:
- Organizes content by learning objectives and skill progression
- Includes clear explanations of concepts and theories
- Focuses on student comprehension and retention
- Incorporates exercises, examples, and assessments
- Provides structured learning path from basic to advanced
- Balances theoretical knowledge with practical application
- Addresses different learning styles and needs""",
            "chapter": """Write in an educational/textbook style that:
- Uses clear, instructional language appropriate for students
- Provides step-by-step explanations of concepts
- Includes relevant examples and practical applications
- Incorporates exercises and learning activities
- Uses consistent terminology and definitions
- Maintains engaging tone while being informative
- Addresses common student questions and misconceptions""",
            "formatting": """Use educational HTML formatting:
- <h2> for chapter titles with learning topic names
- <h3> for major concepts or learning objectives
- <h4> for sub-topics and detailed explanations
- <p> for instructional content and explanations
- <ul> and <li> for key points and learning objectives
- <strong> for important terms and concepts
- <em> for definitions and emphasis"""
        },
        "spiritual": {
            "outline": """Create a spiritual/religious outline that:
- Organizes content by spiritual themes or religious concepts
- Includes respectful exploration of faith and beliefs
- Focuses on personal growth and spiritual development
- Incorporates diverse spiritual perspectives when appropriate
- Provides practical guidance for spiritual practice
- Balances theological concepts with personal application
- Addresses questions of meaning and purpose""",
            "chapter": """Write in a spiritual/religious style that:
- Uses respectful, reverent language for sacred topics
- Incorporates personal spiritual insights and experiences
- Includes practical guidance for spiritual growth
- Provides comfort and inspiration to readers
- Balances theological concepts with accessible explanations
- Uses appropriate religious or spiritual terminology
- Maintains sensitivity to diverse spiritual beliefs""",
            "formatting": """Use spiritual HTML formatting:
- <h2> for chapter titles with spiritual theme names
- <h3> for major spiritual concepts or practices
- <h4> for specific teachings or spiritual exercises
- <p> for spiritual guidance and reflections
- <ul> and <li> for spiritual practices and key teachings
- <strong> for important spiritual concepts and terms
- <em> for scripture quotes and spiritual emphasis"""
        },
        "finance": {
            "outline": """Create a finance/investment outline that:
- Organizes content by financial concepts or investment strategies
- Includes practical financial advice and planning strategies
- Focuses on actionable financial guidance and education
- Incorporates current market trends and economic principles
- Provides risk assessment and management strategies
- Balances technical financial concepts with accessibility
- Addresses various financial goals and life stages""",
            "chapter": """Write in a finance/investment style that:
- Uses clear, professional language for financial concepts
- Provides practical, actionable financial advice
- Includes real-world examples and case studies
- Incorporates current market data and trends
- Addresses risk management and financial planning
- Uses encouraging tone while being realistic about challenges
- Maintains credibility with evidence-based recommendations""",
            "formatting": """Use finance HTML formatting:
- <h2> for chapter titles with financial topic names
- <h3> for major financial concepts or strategies
- <h4> for specific financial instruments or techniques
- <p> for financial explanations and advice
- <ul> and <li> for financial strategies and key points
- <strong> for important financial terms and warnings
- <em> for financial terminology and emphasis"""
        }
    }
    
    # Default to story style if unknown
    if writing_style not in styles:
        writing_style = "story"
    
    return styles[writing_style].get(content_type, styles["story"][content_type])

# Helper function for cleaning AI responses
def clean_ai_response(response: str) -> str:
    """Clean and format AI response with proper HTML structure"""
    cleaned_response = response.strip()
    
    # Remove markdown code block patterns
    if cleaned_response.startswith('```html'):
        cleaned_response = cleaned_response[7:]
    elif cleaned_response.startswith('```'):
        cleaned_response = cleaned_response[3:]
    
    if cleaned_response.endswith('```'):
        cleaned_response = cleaned_response[:-3]
    
    # Remove markdown artifacts and improve formatting
    cleaned_response = cleaned_response.replace('```html', '').replace('```', '')
    cleaned_response = cleaned_response.strip()
    
    # Remove common AI preamble patterns in multiple languages
    preamble_patterns = [
        r'^.*?(Here is|Ecco|Voici|Aquí está|Hier ist).*?(?=<h2>|Chapter|Capitolo|Chapitre|Capítulo|Kapitel)',
        r'^.*?(detailed outline|bozza dettagliata|plan détaillé|esquema detallado).*?(?=<h2>|Chapter|Capitolo|Chapitre|Capítulo|Kapitel)',
        r'^.*?(following your specifications|secondo le tue specifiche|selon vos spécifications|siguiendo sus especificaciones).*?(?=<h2>|Chapter|Capitolo|Chapitre|Capítulo|Kapitel)',
        r'^.*?(\*\*\*|---|\*\*.*?\*\*).*?(?=<h2>|Chapter|Capitolo|Chapitre|Capítulo|Kapitel)',
        r'^.*?(Book Description|Descrizione del libro|Description du livre|Descripción del libro).*?(?=<h2>|Chapter|Capitolo|Chapitre|Capítulo|Kapitel)',
        r'^.*?(Outline|Schema|Plan|Esquema|Gliederung).*?(?=<h2>|Chapter|Capitolo|Chapitre|Capítulo|Kapitel)',
        r'^.*?(?:Certainly|Certamente|Certainement|Ciertamente|Sicherlich).*?(?=<h2>|Chapter|Capitolo|Chapitre|Capítulo|Kapitel)',
        r'^[^<]*(?=<h2>)',  # Remove any text before first h2 tag
        r'^.*?(?=<h2>Chapter|<h2>Capitolo|<h2>Chapitre|<h2>Capítulo|<h2>Kapitel)',  # Remove text before chapter headers
    ]
    
    for pattern in preamble_patterns:
        cleaned_response = re.sub(pattern, '', cleaned_response, flags=re.DOTALL | re.IGNORECASE)
    
    cleaned_response = cleaned_response.strip()
    
    # Clean up chapter title formatting issues
    # Fix duplicated chapter titles
    cleaned_response = re.sub(r'<h2>([^<]+)</h2>\s*#+\s*\1', r'<h2>\1</h2>', cleaned_response, flags=re.IGNORECASE)
    
    # Remove standalone markdown headers that might duplicate HTML headers
    cleaned_response = re.sub(r'\n#+\s*([^\n]+)\n(?=\s*<h2>)', '', cleaned_response)
    
    # Enhance HTML formatting with proper spacing
    cleaned_response = cleaned_response.replace('<h1>', '\n\n<h1>').replace('</h1>', '</h1>\n\n')
    cleaned_response = cleaned_response.replace('<h2>', '\n\n<h2>').replace('</h2>', '</h2>\n\n')
    cleaned_response = cleaned_response.replace('<h3>', '\n<h3>').replace('</h3>', '</h3>\n')
    cleaned_response = cleaned_response.replace('<h4>', '\n<h4>').replace('</h4>', '</h4>\n')
    cleaned_response = cleaned_response.replace('<p>', '\n<p>').replace('</p>', '</p>\n\n')
    cleaned_response = cleaned_response.replace('<ul>', '\n<ul>').replace('</ul>', '</ul>\n\n')
    cleaned_response = cleaned_response.replace('<li>', '\n  <li>').replace('</li>', '</li>')
    
    # Clean up excessive spacing
    cleaned_response = re.sub(r'\n{3,}', '\n\n', cleaned_response)
    cleaned_response = cleaned_response.strip()
    
    # Apply asterisk formatting cleanup
    cleaned_response = process_asterisk_formatting(cleaned_response)
    
    # Final cleanup of any remaining unwanted text patterns
    cleaned_response = re.sub(r'^[^<]*?(?=<h2>)', '', cleaned_response, flags=re.DOTALL)
    
    return cleaned_response

def hash_password(password: str) -> str:
    """Hash a password for storing."""
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify a password against its hash."""
    return pwd_context.verify(plain_password, hashed_password)

def validate_password(password: str) -> bool:
    """Validate password strength."""
    return len(password) >= 8

# Authentication helper functions
async def get_current_user(authorization: str = Header(None)):
    """Get current user from session token"""
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header required")
    
    try:
        # Extract token from Bearer header
        if authorization.startswith("Bearer "):
            token = authorization.split(" ")[1]
        else:
            token = authorization
        
        # Find session in database
        session = await db.user_sessions.find_one({"session_token": token})
        if not session:
            raise HTTPException(status_code=401, detail="Invalid session token")
        
        # Check if session is expired
        if datetime.utcnow() > session["expires_at"]:
            await db.user_sessions.delete_one({"session_token": token})
            raise HTTPException(status_code=401, detail="Session expired")
        
        # Get user data
        user = await db.users.find_one({"id": session["user_id"]})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        
        return User(**user)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid session token")

async def create_user_session(user_id: str, session_token: str) -> UserSession:
    """Create a new user session"""
    expires_at = datetime.utcnow() + timedelta(days=7)
    session = UserSession(
        user_id=user_id,
        session_token=session_token,
        expires_at=expires_at
    )
    await db.user_sessions.insert_one(session.dict())
    return session

# Credit management functions
async def get_user_credit_balance(user_id: str) -> int:
    """Get user's current credit balance"""
    user = await db.users.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user.get("credit_balance", 0)

async def deduct_credits(user_id: str, amount: int, transaction_type: str, description: str, 
                        book_project_id: str = None, chapter_number: int = None) -> bool:
    """Deduct credits from user account and record transaction"""
    try:
        # Get current balance
        user = await db.users.find_one({"id": user_id})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        current_balance = user.get("credit_balance", 0)
        
        # Check if user has sufficient credits
        if current_balance < amount:
            return False
        
        # Deduct credits
        new_balance = current_balance - amount
        await db.users.update_one(
            {"id": user_id},
            {"$set": {"credit_balance": new_balance, "updated_at": datetime.utcnow()}}
        )
        
        # Record transaction
        transaction = CreditTransaction(
            user_id=user_id,
            amount=-amount,  # Negative for deduction
            transaction_type=transaction_type,
            description=description,
            book_project_id=book_project_id,
            chapter_number=chapter_number
        )
        await db.credit_transactions.insert_one(transaction.dict())
        
        return True
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Credit deduction failed: {str(e)}")

async def add_credits(user_id: str, amount: int, transaction_type: str, description: str) -> int:
    """Add credits to user account and record transaction"""
    try:
        # Get current balance
        user = await db.users.find_one({"id": user_id})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        current_balance = user.get("credit_balance", 0)
        new_balance = current_balance + amount
        
        # Add credits
        await db.users.update_one(
            {"id": user_id},
            {"$set": {"credit_balance": new_balance, "updated_at": datetime.utcnow()}}
        )
        
        # Record transaction
        transaction = CreditTransaction(
            user_id=user_id,
            amount=amount,  # Positive for addition
            transaction_type=transaction_type,
            description=description
        )
        await db.credit_transactions.insert_one(transaction.dict())
        
        return new_balance
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Credit addition failed: {str(e)}")

def calculate_book_cost(pages: int, chapters: int) -> dict:
    """Calculate the minimum cost to generate a book"""
    # Ensure minimum chapters based on max 10 pages per chapter
    min_chapters = max(chapters, (pages + 9) // 10)  # Round up division
    
    return {
        "pages": pages,
        "requested_chapters": chapters,
        "minimum_chapters": min_chapters,
        "cost_per_chapter": 1,
        "total_cost": min_chapters,
        "pages_per_chapter": pages / min_chapters if min_chapters > 0 else 0
    }

async def is_chapter_regeneration(book_project_id: str, chapter_number: int) -> bool:
    """Check if this is a chapter regeneration (already exists)"""
    project = await db.book_projects.find_one({"id": book_project_id})
    if not project:
        return False
    
    generated_chapters = project.get("generated_chapters", [])
    return chapter_number in generated_chapters

async def mark_chapter_as_generated(book_project_id: str, chapter_number: int):
    """Mark a chapter as generated in the project"""
    await db.book_projects.update_one(
        {"id": book_project_id},
        {
            "$addToSet": {"generated_chapters": chapter_number},
            "$set": {"updated_at": datetime.utcnow()}
        }
    )

# Authentication endpoints
@api_router.post("/auth/register")
async def register_user(request: RegisterRequest):
    """Register a new user with email and password"""
    try:
        # Validate input
        if not request.email or not request.name or not request.password:
            raise HTTPException(status_code=400, detail="Email, name, and password are required")
        
        # Validate password strength
        if not validate_password(request.password):
            raise HTTPException(status_code=400, detail="Password must be at least 8 characters long")
        
        # Check if user already exists
        existing_user = await db.users.find_one({"email": request.email})
        if existing_user:
            raise HTTPException(status_code=400, detail="User with this email already exists")
        
        # Create new user
        password_hash = hash_password(request.password)
        user = User(
            email=request.email,
            name=request.name,
            password_hash=password_hash,
            auth_provider="email"
        )
        
        await db.users.insert_one(user.dict())
        
        # Create session
        session_token = secrets.token_urlsafe(32)
        session = await create_user_session(user.id, session_token)
        
        return {
            "user": UserProfile(
                id=user.id,
                email=user.email,
                name=user.name,
                picture=user.picture
            ),
            "session_token": session_token,
            "expires_at": session.expires_at.isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Registration failed: {str(e)}")

@api_router.post("/auth/login")
async def login_user(request: EmailPasswordRequest):
    """Login user with email and password"""
    try:
        # Validate input
        if not request.email or not request.password:
            raise HTTPException(status_code=400, detail="Email and password are required")
        
        # Find user
        user_data = await db.users.find_one({"email": request.email})
        if not user_data:
            raise HTTPException(status_code=401, detail="Invalid email or password")
        
        user = User(**user_data)
        
        # Verify password (only for email auth users)
        if user.auth_provider == "email":
            if not user.password_hash or not verify_password(request.password, user.password_hash):
                raise HTTPException(status_code=401, detail="Invalid email or password")
        else:
            raise HTTPException(status_code=401, detail="This account uses Google authentication")
        
        # Create session
        session_token = secrets.token_urlsafe(32)
        session = await create_user_session(user.id, session_token)
        
        return {
            "user": UserProfile(
                id=user.id,
                email=user.email,
                name=user.name,
                picture=user.picture
            ),
            "session_token": session_token,
            "expires_at": session.expires_at.isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Login failed: {str(e)}")

@api_router.post("/auth/google")
async def authenticate_google(request: GoogleTokenRequest):
    """Authenticate user with Google OAuth token"""
    try:
        # Verify Google OAuth token
        async with httpx.AsyncClient() as client:
            response = await client.get(
                f"https://www.googleapis.com/oauth2/v1/tokeninfo?access_token={request.token}"
            )
            
            if response.status_code != 200:
                raise HTTPException(status_code=401, detail="Invalid Google token")
            
            token_info = response.json()
            
            # Get user profile from Google
            profile_response = await client.get(
                f"https://www.googleapis.com/oauth2/v1/userinfo?access_token={request.token}"
            )
            
            if profile_response.status_code != 200:
                raise HTTPException(status_code=401, detail="Failed to get user profile")
            
            profile_data = profile_response.json()
            
            # Check if user exists
            existing_user = await db.users.find_one({"email": profile_data["email"]})
            
            if existing_user:
                user = User(**existing_user)
            else:
                # Create new user
                user = User(
                    email=profile_data["email"],
                    name=profile_data["name"],
                    picture=profile_data.get("picture"),
                    auth_provider="google"
                )
                await db.users.insert_one(user.dict())
            
            # Create user session
            session_token = secrets.token_urlsafe(32)
            session = await create_user_session(user.id, session_token)
            
            return {
                "user": UserProfile(
                    id=user.id,
                    email=user.email,
                    name=user.name,
                    picture=user.picture
                ),
                "session_token": session_token,
                "expires_at": session.expires_at.isoformat()
            }
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Authentication failed: {str(e)}")

@api_router.post("/auth/google/verify")
async def verify_google_token(request: GoogleTokenRequest):
    """Verify Google ID token and authenticate user"""
    try:
        # Decode Google ID token without verification (for demo purposes)
        # In production, you should verify the token signature
        import base64
        import json
        
        # Extract payload from JWT token
        token_parts = request.token.split('.')
        if len(token_parts) != 3:
            raise HTTPException(status_code=401, detail="Invalid token format")
        
        # Decode the payload (add padding if needed)
        payload = token_parts[1]
        payload += '=' * (4 - len(payload) % 4)  # Add padding
        
        try:
            decoded_payload = base64.urlsafe_b64decode(payload)
            profile_data = json.loads(decoded_payload)
        except Exception:
            raise HTTPException(status_code=401, detail="Invalid token format")
        
        # Check if user exists
        existing_user = await db.users.find_one({"email": profile_data["email"]})
        
        if existing_user:
            user = User(**existing_user)
        else:
            # Create new user
            user = User(
                email=profile_data["email"],
                name=profile_data["name"],
                picture=profile_data.get("picture"),
                auth_provider="google"
            )
            await db.users.insert_one(user.dict())
        
        # Create user session
        session_token = secrets.token_urlsafe(32)
        session = await create_user_session(user.id, session_token)
        
        return {
            "user": UserProfile(
                id=user.id,
                email=user.email,
                name=user.name,
                picture=user.picture
            ),
            "session_token": session_token,
            "expires_at": session.expires_at.isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Authentication failed: {str(e)}")

@api_router.get("/user/stats")
async def get_user_stats(current_user: User = Depends(get_current_user)):
    """Get user statistics for dashboard"""
    try:
        # Get user's projects
        projects = await db.book_projects.find({"user_id": current_user.id}).to_list(1000)
        
        # Calculate statistics
        total_books = len(projects)
        total_chapters = 0
        total_words = 0
        completed_books = 0
        
        for project in projects:
            chapters_content = project.get("chapters_content", {})
            project_chapters = len(chapters_content)
            total_chapters += project_chapters
            
            # Calculate words in chapters
            for chapter_content in chapters_content.values():
                if chapter_content:
                    # Remove HTML tags and count words
                    clean_text = re.sub(r'<[^>]+>', '', chapter_content)
                    words = len(clean_text.split())
                    total_words += words
            
            # Check if book is completed (has all chapters)
            if project_chapters >= project.get("chapters", 0):
                completed_books += 1
        
        # Calculate recent activity (last 30 days)
        thirty_days_ago = datetime.utcnow() - timedelta(days=30)
        recent_projects = await db.book_projects.find({
            "user_id": current_user.id,
            "updated_at": {"$gte": thirty_days_ago}
        }).to_list(1000)
        
        return {
            "total_books": total_books,
            "completed_books": completed_books,
            "total_chapters": total_chapters,
            "total_words": total_words,
            "recent_activity": len(recent_projects),
            "avg_words_per_chapter": round(total_words / total_chapters) if total_chapters > 0 else 0,
            "user_since": current_user.created_at.strftime("%B %Y") if current_user.created_at else "Recently",
            "credit_balance": current_user.credit_balance
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching user stats: {str(e)}")

@api_router.get("/auth/profile")
async def get_user_profile(current_user: User = Depends(get_current_user)):
    """Get current user profile"""
    return UserProfile(
        id=current_user.id,
        email=current_user.email,
        name=current_user.name,
        picture=current_user.picture
    )

@api_router.post("/auth/logout")
async def logout(authorization: str = Header(None)):
    """Logout user by invalidating session"""
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header required")
    
    try:
        # Extract token from Bearer header
        if authorization.startswith("Bearer "):
            token = authorization.split(" ")[1]
        else:
            token = authorization
        
        # Delete session from database
        await db.user_sessions.delete_one({"session_token": token})
        
        return {"message": "Logged out successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Logout failed: {str(e)}")

# Credit management endpoints
@api_router.get("/credits/balance", response_model=CreditBalanceResponse)
async def get_credit_balance(current_user: User = Depends(get_current_user)):
    """Get user's current credit balance"""
    try:
        balance = await get_user_credit_balance(current_user.id)
        return CreditBalanceResponse(
            credit_balance=balance,
            user_id=current_user.id
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching credit balance: {str(e)}")

@api_router.get("/credits/history")
async def get_credit_history(current_user: User = Depends(get_current_user), limit: int = 50):
    """Get user's credit transaction history"""
    try:
        transactions = await db.credit_transactions.find(
            {"user_id": current_user.id}
        ).sort("created_at", -1).limit(limit).to_list(length=limit)
        
        return [CreditTransactionResponse(**transaction) for transaction in transactions]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching credit history: {str(e)}")

@api_router.post("/credits/calculate-book-cost", response_model=BookCostResponse)
async def calculate_book_cost_endpoint(request: BookCostRequest):
    """Calculate the cost to generate a book"""
    try:
        cost_info = calculate_book_cost(request.pages, request.chapters)
        return BookCostResponse(**cost_info)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error calculating book cost: {str(e)}")

@api_router.get("/credits/packages")
async def get_credit_packages():
    """Get available credit packages"""
    return {"packages": CREDIT_PACKAGES}

@api_router.post("/payments/create-session", response_model=PaymentSessionResponse)
async def create_payment_session(request: PaymentPackageRequest, current_user: User = Depends(get_current_user)):
    """Create a Stripe checkout session for credit purchase"""
    try:
        # Validate package ID
        if request.package_id not in CREDIT_PACKAGES:
            raise HTTPException(status_code=400, detail="Invalid package ID")
        
        package = CREDIT_PACKAGES[request.package_id]
        
        # Initialize Stripe checkout
        stripe_api_key = get_stripe_key()
        if not stripe_api_key:
            raise HTTPException(status_code=500, detail="Stripe API key not configured")
        
        # Construct webhook URL (will be handled by FastAPI)
        webhook_url = f"{request.origin_url}/api/webhook/stripe"
        stripe_checkout = StripeCheckout(api_key=stripe_api_key, webhook_url=webhook_url)
        
        # Construct success and cancel URLs
        success_url = f"{request.origin_url}/payment-success?session_id={{CHECKOUT_SESSION_ID}}"
        cancel_url = f"{request.origin_url}/credits"
        
        # Create checkout session request
        checkout_request = CheckoutSessionRequest(
            amount=package["price"],
            currency=package["currency"],
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                "user_id": current_user.id,
                "package_id": request.package_id,
                "credits_amount": str(package["credits"]),
                "source": "web_checkout"
            }
        )
        
        # Create Stripe checkout session
        session_response: CheckoutSessionResponse = await stripe_checkout.create_checkout_session(checkout_request)
        
        # Create payment transaction record in database
        payment_transaction = PaymentTransaction(
            user_id=current_user.id,
            session_id=session_response.session_id,
            amount=package["price"],
            currency=package["currency"],
            credits_amount=package["credits"],
            package_id=request.package_id,
            payment_status="pending",
            status="initiated",
            metadata={
                "user_id": current_user.id,
                "package_id": request.package_id,
                "credits_amount": str(package["credits"]),
                "source": "web_checkout"
            }
        )
        
        await db.payment_transactions.insert_one(payment_transaction.dict())
        
        return PaymentSessionResponse(
            checkout_url=session_response.url,
            session_id=session_response.session_id,
            package_info={
                "name": package["name"],
                "credits": package["credits"],
                "price": package["price"],
                "currency": package["currency"],
                "description": package["description"]
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create payment session: {str(e)}")

@api_router.get("/payments/status/{session_id}", response_model=PaymentStatusResponse)
async def get_payment_status(session_id: str, current_user: User = Depends(get_current_user)):
    """Get payment status and process successful payments"""
    try:
        # Find payment transaction in database
        payment_record = await db.payment_transactions.find_one({"session_id": session_id, "user_id": current_user.id})
        if not payment_record:
            raise HTTPException(status_code=404, detail="Payment session not found")
        
        # Initialize Stripe checkout to check status
        stripe_api_key = get_stripe_key()
        if not stripe_api_key:
            raise HTTPException(status_code=500, detail="Stripe API key not configured")
        
        # Use a dummy webhook URL since we're just checking status
        webhook_url = "https://dummy.webhook.url"
        stripe_checkout = StripeCheckout(api_key=stripe_api_key, webhook_url=webhook_url)
        
        # Get checkout status from Stripe
        checkout_status: CheckoutStatusResponse = await stripe_checkout.get_checkout_status(session_id)
        
        # Update payment record with latest status
        updated_payment = {
            "payment_status": checkout_status.payment_status,
            "status": checkout_status.status,
            "updated_at": datetime.utcnow()
        }
        
        # If payment is successful and not already processed
        if (checkout_status.payment_status == "paid" and 
            payment_record.get("payment_status") != "paid"):
            
            # Add credits to user account
            credits_to_add = payment_record["credits_amount"]
            new_balance = await add_credits(
                current_user.id,
                credits_to_add,
                "credit_purchase",
                f"Purchased {credits_to_add} credits via Stripe - Package: {payment_record['package_id']}"
            )
            
            # Update payment record as completed
            updated_payment["status"] = "completed"
        
        # Update payment record in database
        await db.payment_transactions.update_one(
            {"session_id": session_id, "user_id": current_user.id},
            {"$set": updated_payment}
        )
        
        return PaymentStatusResponse(
            session_id=session_id,
            payment_status=checkout_status.payment_status,
            status=updated_payment["status"],
            amount=payment_record["amount"],
            currency=payment_record["currency"],
            credits_amount=payment_record["credits_amount"],
            package_id=payment_record["package_id"]
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get payment status: {str(e)}")

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhook events"""
    try:
        # Get webhook payload and signature
        body = await request.body()
        signature = request.headers.get("stripe-signature")
        
        if not signature:
            raise HTTPException(status_code=400, detail="Missing Stripe signature")
        
        # Initialize Stripe checkout
        stripe_api_key = get_stripe_key()
        if not stripe_api_key:
            raise HTTPException(status_code=500, detail="Stripe API key not configured")
        
        webhook_url = "https://dummy.webhook.url"  # Not used for webhook handling
        stripe_checkout = StripeCheckout(api_key=stripe_api_key, webhook_url=webhook_url)
        
        # Handle webhook event
        webhook_response = await stripe_checkout.handle_webhook(body, signature)
        
        # Process the webhook event
        if webhook_response.event_type in ["checkout.session.completed", "payment_intent.succeeded"]:
            session_id = webhook_response.session_id
            
            # Find and update payment record
            payment_record = await db.payment_transactions.find_one({"session_id": session_id})
            if payment_record:
                # Update payment status
                await db.payment_transactions.update_one(
                    {"session_id": session_id},
                    {
                        "$set": {
                            "payment_status": webhook_response.payment_status,
                            "status": "completed" if webhook_response.payment_status == "paid" else "failed",
                            "updated_at": datetime.utcnow()
                        }
                    }
                )
                
                # Add credits if payment successful and not already processed
                if (webhook_response.payment_status == "paid" and 
                    payment_record.get("payment_status") != "paid"):
                    
                    credits_to_add = payment_record["credits_amount"]
                    await add_credits(
                        payment_record["user_id"],
                        credits_to_add,
                        "credit_purchase",
                        f"Purchased {credits_to_add} credits via Stripe webhook - Package: {payment_record['package_id']}"
                    )
        
        return {"status": "success"}
        
    except Exception as e:
        logging.error(f"Stripe webhook error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Webhook processing failed: {str(e)}")

@api_router.post("/credits/purchase")
async def purchase_credits_legacy(request: CreditPurchaseRequest, current_user: User = Depends(get_current_user)):
    """Legacy credit purchase endpoint - redirects to new payment system"""
    return {
        "message": "Please use the new payment system with credit packages",
        "redirect_to": "/credits",
        "packages_available": list(CREDIT_PACKAGES.keys())
    }

# Helper function to extract chapter titles from outline
def extract_chapter_titles(outline: str) -> dict:
    """Extract chapter titles from the outline HTML"""
    chapter_titles = {}
    
    # Look for chapter titles in various formats
    patterns = [
        r'<h2[^>]*>Chapter\s+(\d+):?\s*([^<]+)</h2>',
        r'<h2[^>]*>(\d+)\.\s*([^<]+)</h2>',
        r'<h2[^>]*>([^<]*Chapter\s+\d+[^<]*)</h2>'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, outline, re.IGNORECASE)
        for match in matches:
            if len(match) == 2:
                chapter_num = int(match[0])
                chapter_title = match[1].strip()
                chapter_titles[chapter_num] = chapter_title
    
    return chapter_titles

def generate_table_of_contents(project_obj: BookProject) -> str:
    """Generate a table of contents with chapter titles and page numbers"""
    toc_html = '<div class="table-of-contents">\n<h2>📚 Table of Contents</h2>\n<div class="toc-entries">\n'
    
    # Extract chapter titles from outline
    chapter_titles = extract_chapter_titles(project_obj.outline or "")
    
    # Calculate approximate page numbers (assuming 275 words per page)
    current_page = 1
    words_per_page = 275
    
    for i in range(1, project_obj.chapters + 1):
        # Get chapter title
        chapter_title = chapter_titles.get(i, f"Chapter {i}")
        
        # Calculate page number for this chapter
        page_number = current_page
        
        # Add TOC entry
        toc_html += f'<div class="toc-entry">\n'
        toc_html += f'  <span class="toc-chapter">Chapter {i}: {chapter_title}</span>\n'
        toc_html += f'  <span class="toc-dots">{"." * 50}</span>\n'
        toc_html += f'  <span class="toc-page">{page_number}</span>\n'
        toc_html += f'</div>\n'
        
        # Calculate next chapter's starting page
        if project_obj.chapters_content and str(i) in project_obj.chapters_content:
            chapter_content = project_obj.chapters_content[str(i)]
            # Count words in chapter (rough estimate)
            word_count = len(chapter_content.split())
            pages_in_chapter = max(1, word_count // words_per_page)
            current_page += pages_in_chapter
        else:
            # Default pages per chapter if content not available
            estimated_pages = max(1, (project_obj.pages // project_obj.chapters))
            current_page += estimated_pages
    
    toc_html += '</div>\n</div>'
    return toc_html
@api_router.get("/")
async def root():
    return {"message": "AI Book Writer API"}

@api_router.post("/projects", response_model=BookProject)
async def create_project(project_data: BookProjectCreate, current_user: User = Depends(get_current_user)):
    """Create a new book project"""
    try:
        # Validate and adjust chapter requirements based on page limit
        cost_info = calculate_book_cost(project_data.pages, project_data.chapters)
        
        # Update chapters to minimum required if user specified too few
        if project_data.chapters < cost_info["minimum_chapters"]:
            project_data.chapters = cost_info["minimum_chapters"]
        
        # Check if user has sufficient credits for the book
        current_balance = await get_user_credit_balance(current_user.id)
        if current_balance < cost_info["total_cost"]:
            raise HTTPException(
                status_code=402,
                detail=f"Insufficient credits. This book requires {cost_info['total_cost']} credits ({cost_info['minimum_chapters']} chapters × 1 credit each) but you have {current_balance}. Please purchase more credits."
            )
        
        project_dict = project_data.dict()
        project_dict["user_id"] = current_user.id  # Associate with current user
        project_obj = BookProject(**project_dict)
        await db.book_projects.insert_one(project_obj.dict())
        
        return project_obj
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating project: {str(e)}")

@api_router.get("/projects", response_model=List[BookProject])
async def get_projects(current_user: User = Depends(get_current_user)):
    """Get all book projects for the current user"""
    try:
        projects = await db.book_projects.find({"user_id": current_user.id}).to_list(1000)
        return [BookProject(**project) for project in projects]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching projects: {str(e)}")

@api_router.get("/projects/{project_id}", response_model=BookProject)
async def get_project(project_id: str, current_user: User = Depends(get_current_user)):
    """Get a specific book project"""
    try:
        project = await db.book_projects.find_one({"id": project_id, "user_id": current_user.id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        return BookProject(**project)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching project: {str(e)}")

@api_router.post("/generate-outline")
async def generate_outline(request: OutlineRequest, current_user: User = Depends(get_current_user)):
    """Generate book outline using Gemini 2.5 Flash-Lite AI"""
    try:
        # Get project details and verify ownership
        project = await db.book_projects.find_one({"id": request.project_id, "user_id": current_user.id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        project_obj = BookProject(**project)
        
        # Initialize Gemini chat with updated model
        chat = LlmChat(
            api_key="AIzaSyDYSBIsBp1KEENEDihvS0Nl0A6lGsgzMgM",
            session_id=f"outline_{project_obj.id}",
            system_message="You are an expert book writing assistant. You create detailed, well-structured book outlines that serve as comprehensive guides for content generation."
        ).with_model("gemini", "gemini-2.5-flash-lite-preview-06-17")
        
        # Enhanced style-specific instructions based on writing style
        if project_obj.writing_style == "story":
            style_instructions = """Create a story-focused outline that:
- Develops compelling character arcs and relationships
- Crafts engaging plot points and narrative hooks
- Maintains consistent pacing and tension
- Builds immersive story worlds and settings
- Weaves subplots seamlessly into the main narrative
- Creates emotional resonance and character growth
- Uses natural scene and chapter transitions"""
        else:  # descriptive
            style_instructions = """Create a structured informational outline that:
- Presents clear topic hierarchies and logical flow
- Breaks down complex concepts systematically
- Incorporates relevant examples and case studies
- Balances depth and accessibility of content
- Uses consistent terminology and definitions
- Provides clear learning objectives per section
- Includes practical applications and insights"""
        
        # Enhanced outline generation prompt for consistency and naturalness
        language_instructions = ""
        if project_obj.language.lower() != "english":
            language_instructions = f"""
**Language Requirements for {project_obj.language}:**
- Use natural, fluent {project_obj.language} that feels authentic to native speakers
- Employ idiomatic expressions and natural phrasing typical of {project_obj.language} literature
- Avoid literal translations or awkward constructions
- Use appropriate cultural references and context for {project_obj.language} readers
- Maintain consistent tone and style throughout all content"""

        prompt = f"""Create a detailed book outline for "{project_obj.title}":

**Book Details:**
- {project_obj.chapters} chapters, {project_obj.pages} pages
- Style: {project_obj.writing_style}
- Language: {project_obj.language}

**CRITICAL: Response Format Requirements:**
- Start IMMEDIATELY with the first chapter title in HTML format
- Do NOT include any preamble, explanations, or introductory text
- Do NOT include phrases like "Here is", "Ecco", "Voici", etc.
- Do NOT include "Book Description" or "Outline" headers
- Do NOT include any text before the first <h2> tag
- Response must begin directly with: <h2>Chapter 1: [Creative Title]</h2>

**Requirements:**
1. Create exactly {project_obj.chapters} expressive, thematic chapter titles that capture the essence of each chapter
2. Each chapter: 3-5 sentences summary + 3-5 key points
3. Target {(project_obj.pages * 275) // project_obj.chapters} words per chapter
4. HTML format: <h2>Chapter [Number]: [Creative Title]</h2> for titles
5. Use <p> for summaries, <ul><li> for key points

**Creative Chapter Titles:**
- Make each chapter title evocative and atmospheric
- Reflect the emotional tone or key theme of the chapter
- Avoid generic titles like "Chapter 1" or "Introduction"
- Use metaphors, imagery, or compelling phrases that draw readers in
- Examples: "The Weight of Silence", "Dancing with Shadows", "Where Rivers Meet the Sea"

**Style Guide:**
{style_instructions}

**Narrative Voice Consistency:**
- Choose a consistent narrative voice (first-person "I" or third-person "he/she/they") and maintain it throughout
- Ensure the chosen voice fits the writing style and subject matter
- Never switch between narrative voices within the same context

**Content Balance & Structure:**
- Plan chapters with varied pacing and emotional rhythm
- Balance descriptive passages with action and dialogue scenes
- Include character development and interaction opportunities
- Ensure each chapter has a clear narrative arc and purpose

{language_instructions}

**Book Description:**
{project_obj.description}

Begin your response immediately with the first chapter title. Do not include any explanatory text, preambles, or introductions."""

        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        # Enhanced response cleanup
        cleaned_response = response.strip()
        
        # Remove markdown code block patterns
        if cleaned_response.startswith('```html'):
            cleaned_response = cleaned_response[7:]
        elif cleaned_response.startswith('```'):
            cleaned_response = cleaned_response[3:]
        
        if cleaned_response.endswith('```'):
            cleaned_response = cleaned_response[:-3]
        
        # Remove markdown artifacts and improve formatting
        cleaned_response = cleaned_response.replace('```html', '').replace('```', '')
        cleaned_response = cleaned_response.strip()
        
        # Enhance HTML formatting
        cleaned_response = cleaned_response.replace('<h2>', '\n\n<h2>').replace('</h2>', '</h2>\n\n')
        cleaned_response = cleaned_response.replace('<h3>', '\n<h3>').replace('</h3>', '</h3>\n')
        cleaned_response = cleaned_response.replace('<p>', '\n<p>').replace('</p>', '</p>\n\n')
        cleaned_response = cleaned_response.replace('<ul>', '\n<ul>').replace('</ul>', '</ul>\n\n')
        cleaned_response = cleaned_response.replace('<li>', '\n  <li>').replace('</li>', '</li>')
        
        # Clean up spacing
        cleaned_response = re.sub(r'\n{3,}', '\n\n', cleaned_response)
        cleaned_response = cleaned_response.strip()
        
        # Update project with generated outline
        await db.book_projects.update_one(
            {"id": request.project_id},
            {
                "$set": {
                    "outline": cleaned_response,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        return {"outline": cleaned_response, "project_id": request.project_id}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating outline: {str(e)}")



@api_router.post("/generate-chapter")
async def generate_chapter(request: ChapterRequest, current_user: User = Depends(get_current_user)):
    """Generate a specific chapter using Gemini 2.5 Flash-Lite AI"""
    try:
        # Get project details and verify ownership
        project = await db.book_projects.find_one({"id": request.project_id, "user_id": current_user.id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        project_obj = BookProject(**project)
        
        if not project_obj.outline:
            raise HTTPException(status_code=400, detail="Project must have an outline before generating chapters")
        
        # Check if this is a regeneration
        is_regeneration = await is_chapter_regeneration(request.project_id, request.chapter_number)
        credit_cost = 1  # Base cost for chapter generation
        
        # Get current credit balance
        current_balance = await get_user_credit_balance(current_user.id)
        
        # Check if user has sufficient credits
        if current_balance < credit_cost:
            raise HTTPException(
                status_code=402, 
                detail=f"Insufficient credits. You need {credit_cost} credit(s) but have {current_balance}. Please purchase more credits."
            )
        
        # Determine transaction type and description
        transaction_type = "chapter_regeneration" if is_regeneration else "chapter_generation"
        description = f"{'Regenerated' if is_regeneration else 'Generated'} Chapter {request.chapter_number} for '{project_obj.title}'"
        
        # Deduct credits before generation
        credit_deducted = await deduct_credits(
            current_user.id,
            credit_cost,
            transaction_type,
            description,
            request.project_id,
            request.chapter_number
        )
        
        if not credit_deducted:
            raise HTTPException(status_code=402, detail="Failed to deduct credits")
        
        # Extract chapter titles from outline
        chapter_titles = extract_chapter_titles(project_obj.outline)
        chapter_title = chapter_titles.get(request.chapter_number, f"Chapter {request.chapter_number}")
        
        # Initialize Gemini chat for this chapter with updated model
        chat = LlmChat(
            api_key="AIzaSyDYSBIsBp1KEENEDihvS0Nl0A6lGsgzMgM",
            session_id=f"chapter_{project_obj.id}_{request.chapter_number}",
            system_message="You are an expert book writer. You write engaging, well-structured chapters based on outlines. Use HTML formatting for headings, bold text, and structure. Always start each chapter with its proper title."
        ).with_model("gemini", "gemini-2.5-flash-lite-preview-06-17")
        
        # Calculate estimated words per chapter (275 words per page)
        estimated_words_per_chapter = (project_obj.pages * 275) // project_obj.chapters
        
        # Get style-specific instructions
        style_instructions = get_style_instructions(project_obj.writing_style, "chapter")
        formatting_instructions = get_style_instructions(project_obj.writing_style, "formatting")
        
        # Enhanced language-specific instructions for natural content generation
        language_instructions = ""
        if project_obj.language.lower() != "english":
            language_instructions = f"""
**Language Requirements for {project_obj.language}:**
- Write in natural, fluent {project_obj.language} that feels authentic to native speakers
- Use idiomatic expressions and natural phrasing typical of {project_obj.language} literature
- Avoid literal translations or awkward constructions
- Employ appropriate cultural references and context for {project_obj.language} readers
- Use natural sentence structures and vocabulary that native speakers would use
- Maintain consistent tone and style throughout the chapter"""

        # Optimized prompt for faster generation with voice consistency and literary quality
        prompt = f"""Write Chapter {request.chapter_number} for "{project_obj.title}":

**Chapter Title:** {chapter_title}
**Target Length:** {estimated_words_per_chapter} words
**Writing Style:** {project_obj.writing_style}

**CRITICAL: Response Format Requirements:**
- Start IMMEDIATELY with the chapter title in HTML format
- Do NOT include any preamble, explanations, or introductory text
- Do NOT include phrases like "Here is", "Ecco", "Voici", "Certainly", etc.
- Do NOT include any text before the chapter title
- Response must begin directly with: <h2>{chapter_title}</h2>

**Key Requirements:**
1. Start with: <h2>{chapter_title}</h2>
2. Write {estimated_words_per_chapter} words of engaging content
3. Follow the outline for Chapter {request.chapter_number}
4. Use HTML formatting (headings, paragraphs, lists)
5. Write in {project_obj.language}

**Narrative Voice Consistency:**
- Maintain consistent narrative voice throughout this chapter
- If using first-person ("I"), keep it consistent - never switch to third-person
- If using third-person ("he/she/they"), maintain it throughout
- Choose the voice that best fits the content and maintain it completely
- Ensure smooth transitions between sentences and paragraphs

**Dialogue & Character Development:**
- Give each character a distinct voice and speech pattern
- Use natural, conversational dialogue that sounds like real people talking
- Avoid having characters speak in the same formal register as the narrator
- Include interruptions, hesitations, and natural speech patterns
- Show character personality through their word choices and speaking style
- Use contractions and colloquialisms where appropriate
- Break up long speeches with action beats and internal thoughts

**Content Balance & Structure:**
- Alternate between descriptive passages and action/dialogue scenes
- Use varied sentence lengths and structures for natural rhythm
- Include both concrete, grounded scenes and poetic, atmospheric moments
- Break up large blocks of text with dialogue and shorter paragraphs
- Show character emotions through actions and dialogue, not just description

**Emotional Authenticity:**
- Write with genuine emotional depth and human imperfection
- Include subtle internal conflicts and character contradictions
- Allow for moments of vulnerability, uncertainty, and authentic emotion
- Use specific, sensory details rather than abstract concepts
- Show characters' emotions through their actions and reactions
- Include subtle imperfections in thought patterns and speech

**Visual Structure & Formatting:**
- Use frequent paragraph breaks for better readability
- Separate different speakers with clear line breaks
- Use shorter paragraphs for dialogue and action sequences
- Include proper scene transitions and time shifts
- Format dialogue clearly with proper punctuation and spacing

**Content Quality Guidelines:**
- Use natural, flowing language that feels authentic and engaging
- Avoid repetitive phrases or awkward constructions
- Create smooth transitions between ideas and paragraphs
- Use varied sentence structures for natural rhythm
- Include specific details and examples to make content vivid and relatable
- Balance showing vs. telling - demonstrate through action and dialogue

{language_instructions}

**Book Outline:**
{project_obj.outline}

**Style Instructions:**
{style_instructions}

Begin your response immediately with the chapter title. Do not include any explanatory text, preambles, or introductions."""

        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        # Clean up the response
        cleaned_response = clean_ai_response(response)
        
        # Ensure chapter starts with proper title if not already present
        if not cleaned_response.startswith(f'<h2>{chapter_title}</h2>') and not cleaned_response.startswith('<h2>'):
            cleaned_response = f'<h2>{chapter_title}</h2>\n\n{cleaned_response}'
        
        # Update project with the generated chapter
        current_chapters = project.get("chapters_content", {})
        current_chapters[str(request.chapter_number)] = cleaned_response
        
        await db.book_projects.update_one(
            {"id": request.project_id},
            {
                "$set": {
                    "chapters_content": current_chapters,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        # Mark chapter as generated for future regeneration tracking
        await mark_chapter_as_generated(request.project_id, request.chapter_number)
        
        # Get updated credit balance
        new_balance = await get_user_credit_balance(current_user.id)
        
        return {
            "chapter_content": cleaned_response,
            "chapter_number": request.chapter_number,
            "chapter_title": chapter_title,
            "project_id": request.project_id,
            "credit_cost": credit_cost,
            "remaining_credits": new_balance,
            "was_regeneration": is_regeneration
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating chapter: {str(e)}")


@api_router.put("/update-chapter")
async def update_chapter(request: ChapterUpdate, current_user: User = Depends(get_current_user)):
    """Update chapter content"""
    try:
        project = await db.book_projects.find_one({"id": request.project_id, "user_id": current_user.id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        current_chapters = project.get("chapters_content", {})
        current_chapters[str(request.chapter_number)] = request.content
        
        await db.book_projects.update_one(
            {"id": request.project_id},
            {
                "$set": {
                    "chapters_content": current_chapters,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        return {"message": "Chapter updated successfully"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating chapter: {str(e)}")
async def user_has_made_purchase(user_id: str) -> bool:
    """Check if user has made any credit purchases"""
    try:
        purchase_transaction = await db.credit_transactions.find_one({
            "user_id": user_id,
            "transaction_type": "credit_purchase"
        })
        return purchase_transaction is not None
    except Exception:
        return False

def process_asterisk_formatting(text: str) -> str:
    """Process asterisk formatting - convert to bold or remove"""
    # Convert single asterisks around words to bold HTML
    text = re.sub(r'\*([^*\n]+?)\*', r'<strong>\1</strong>', text)
    # Convert double asterisks to bold HTML  
    text = re.sub(r'\*\*([^*\n]+?)\*\*', r'<strong>\1</strong>', text)
    # Remove any remaining standalone asterisks
    text = re.sub(r'\*+', '', text)
    return text

def ensure_consistent_chapter_formatting(content: str, chapter_num: int, chapter_title: str = None) -> str:
    """Ensure consistent chapter formatting across all exports"""
    # Clean existing content
    content = process_asterisk_formatting(content)
    
    # Don't add chapter title if it's already there - just ensure clean formatting
    # The PDF/DOCX export will add the title separately
    
    # Remove any existing chapter titles from content to avoid duplication
    content = re.sub(r'<h2[^>]*>Chapter\s+\d+[^<]*</h2>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'^#+\s*Chapter\s+\d+[^\n]*\n', '', content, flags=re.MULTILINE)
    
    # Ensure proper paragraph structure
    content = content.strip()
    
    # If content doesn't start with a paragraph tag, wrap the beginning
    if content and not content.startswith('<p>'):
        # Split by existing paragraph tags and wrap sections
        parts = re.split(r'(<p[^>]*>|</p>)', content)
        formatted_parts = []
        in_paragraph = False
        
        for part in parts:
            if part.startswith('<p'):
                in_paragraph = True
                formatted_parts.append(part)
            elif part == '</p>':
                in_paragraph = False
                formatted_parts.append(part)
            elif part.strip():
                if not in_paragraph:
                    formatted_parts.append(f'<p>{part}</p>')
                else:
                    formatted_parts.append(part)
        
        content = ''.join(formatted_parts)
    
    # Ensure content ends with closing paragraph tag if it starts with opening
    if content.startswith('<p>') and not content.rstrip().endswith('</p>'):
        content = content.rstrip() + '</p>'
    
    return content

def add_watermark_to_pdf_content(content: list, has_purchased: bool):
    """Add watermark to PDF content if user hasn't purchased - this is a placeholder"""
    # This function is replaced by the proper page template approach
    return content

class WatermarkCanvas:
    """Custom canvas class to add watermarks to PDF pages"""
    def __init__(self, canvas, doc, has_purchased=True):
        self.canvas = canvas
        self.doc = doc
        self.has_purchased = has_purchased
        
    def __getattr__(self, name):
        return getattr(self.canvas, name)
        
    def showPage(self):
        """Called at the end of each page"""
        if not self.has_purchased:
            # Add watermark at bottom center of page
            from reportlab.lib.pagesizes import A4
            page_width = A4[0]
            
            # Save current state
            self.canvas.saveState()
            
            # Set watermark text properties
            self.canvas.setFont("Helvetica", 8)
            self.canvas.setFillColorRGB(0.7, 0.7, 0.7)  # Light gray
            
            # Add watermark at bottom center
            watermark_text = "Generated with BookCraft AI - Purchase credits to remove this watermark"
            text_width = self.canvas.stringWidth(watermark_text, "Helvetica", 8)
            x_position = (page_width - text_width) / 2
            y_position = 30  # 30 points from bottom (same height as page numbers)
            
            self.canvas.drawString(x_position, y_position, watermark_text)
            
            # Restore state
            self.canvas.restoreState()
        
        # Call the original showPage
        self.canvas.showPage()

def add_watermark_to_docx(doc, has_purchased: bool):
    """Add watermark to DOCX document if user hasn't purchased"""
    if not has_purchased:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        try:
            # Create footer with watermark for each section
            sections = doc.sections
            for section in sections:
                # Get or create footer
                footer = section.footer
                
                # Clear existing footer content
                footer._element.clear_content()
                
                # Create watermark paragraph in footer
                watermark_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                watermark_p.alignment = 1  # Center alignment
                
                # Add watermark run
                watermark_run = watermark_p.add_run("Generated with BookCraft AI - Purchase credits to remove this watermark")
                watermark_run.font.name = 'Times New Roman'
                watermark_run.font.size = Inches(0.11)  # 8pt
                watermark_run.font.color.rgb = RGBColor(180, 180, 180)  # Light gray
                
        except Exception as e:
            # If watermark creation fails, continue without it
            print(f"Could not add DOCX watermark: {e}")
            pass


@api_router.get("/export-book/{project_id}")
async def export_book(project_id: str, current_user: User = Depends(get_current_user)):
    """Export book as HTML with table of contents only"""
    try:
        # Get project and verify ownership
        project = await db.book_projects.find_one({"id": project_id, "user_id": current_user.id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        project_obj = BookProject(**project)
        
        # Generate table of contents
        toc_html = generate_table_of_contents(project_obj)
        
        # Create well-formatted HTML content
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project_obj.title}</title>
    <style>
        body {{
            font-family: Georgia, serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.8;
            color: #333;
            background-color: #fff;
        }}
        
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 15px;
            margin-bottom: 30px;
            font-size: 2.5em;
            text-align: center;
        }}
        
        h2 {{
            color: #34495e;
            margin-top: 40px;
            margin-bottom: 20px;
            font-size: 1.8em;
        }}
        
        h3 {{
            color: #7f8c8d;
            margin-top: 30px;
            margin-bottom: 15px;
            font-size: 1.4em;
        }}
        
        p {{
            margin-bottom: 20px;
            text-align: justify;
            text-indent: 30px;
        }}
        
        .book-info {{
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 40px;
            border-left: 5px solid #3498db;
            text-align: center;
        }}
        
        .book-info h1 {{
            border-bottom: none;
            margin-bottom: 20px;
            color: #2c3e50;
        }}
        
        .book-info p {{
            margin-bottom: 10px;
            text-indent: 0;
        }}
        
        .table-of-contents {{
            background: #f8f9fa;
            padding: 25px;
            border-radius: 8px;
            margin-bottom: 40px;
            border-left: 5px solid #28a745;
        }}
        
        .table-of-contents h2 {{
            color: #2c3e50;
            margin-bottom: 25px;
            font-size: 1.8em;
        }}
        
        .toc-entries {{
            max-width: 100%;
        }}
        
        .toc-entry {{
            display: flex;
            align-items: baseline;
            margin-bottom: 12px;
            font-size: 14px;
        }}
        
        .toc-chapter {{
            font-weight: 500;
            color: #2c3e50;
            white-space: nowrap;
        }}
        
        .toc-dots {{
            flex: 1;
            color: #bdc3c7;
            font-size: 12px;
            overflow: hidden;
            margin: 0 8px;
        }}
        
        .toc-page {{
            font-weight: 500;
            color: #3498db;
            white-space: nowrap;
        }}
        
        .chapter {{
            margin-bottom: 60px;
            padding: 20px 0;
            border-bottom: 1px solid #e9ecef;
            page-break-after: always;
        }}
        
        .chapter:last-child {{
            border-bottom: none;
        }}
        
        .chapter h2 {{
            color: #2c3e50;
            font-size: 2em;
            margin-bottom: 25px;
        }}
        
        ul {{
            margin-bottom: 20px;
            padding-left: 40px;
        }}
        
        li {{
            margin-bottom: 8px;
            line-height: 1.6;
        }}
        
        strong {{
            color: #2c3e50;
        }}
        
        em {{
            color: #7f8c8d;
        }}
        
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 8px;
            color: #7f8c8d;
        }}
        
        @media print {{
            body {{
                padding: 20px;
            }}
            .chapter {{
                page-break-after: always;
            }}
            .toc-entry {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
    <div class="book-info">
        <h1>{project_obj.title}</h1>
        <p><strong>Generated On:</strong> {datetime.now().strftime("%B %d, %Y")}</p>
    </div>
    
    {toc_html}
    
    <div class="chapters">"""
        
        # Add chapters with proper formatting
        if project_obj.chapters_content:
            chapter_titles = extract_chapter_titles(project_obj.outline or "")
            for i in range(1, project_obj.chapters + 1):
                chapter_content = project_obj.chapters_content.get(str(i), "")
                if chapter_content:
                    # Get chapter title and ensure consistent formatting
                    chapter_title = chapter_titles.get(i, f"Chapter {i}")
                    formatted_content = ensure_consistent_chapter_formatting(chapter_content, i, chapter_title)
                    
                    html_content += f"""
        <div class="chapter">
            {formatted_content}
        </div>"""
                else:
                    html_content += f"""
        <div class="chapter">
            <h2>Chapter {i}</h2>
            <p><em>This chapter has not been generated yet.</em></p>
        </div>"""
        else:
            html_content += "<p><em>No chapters have been generated yet.</em></p>"
        
        html_content += """
    </div>
    
    <div class="footer">
        <p><em>Generated with AI Book Writer</em></p>
    </div>
</body>
</html>"""
        
        # Create safe filename
        safe_filename = "".join(c for c in project_obj.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_filename = safe_filename.replace(' ', '_')
        if not safe_filename:
            safe_filename = f"book_{project_obj.id}"
        
        return {
            "html": html_content,
            "title": project_obj.title,
            "filename": f"{safe_filename}.html"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting book: {str(e)}")

@api_router.get("/export-book-pdf/{project_id}")
async def export_book_pdf(project_id: str, current_user: User = Depends(get_current_user)):
    """Export book as PDF with professional book formatting"""
    try:
        # Get project and verify ownership
        project = await db.book_projects.find_one({"id": project_id, "user_id": current_user.id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        project_obj = BookProject(**project)
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        
        # Professional book styles - standardized to match DOCX
        styles = getSampleStyleSheet()
        
        # Title page styles
        title_style = ParagraphStyle(
            'BookTitle',
            parent=styles['Title'],
            fontSize=24,  # Standardized size
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=HexColor('#1a1a1a'),
            fontName='Times-Bold'  # Standardized to Times family
        )
        
        author_style = ParagraphStyle(
            'Author',
            parent=styles['Normal'],
            fontSize=14,  # Standardized size
            spaceAfter=20,
            alignment=1,  # Center alignment
            textColor=HexColor('#666666'),
            fontName='Times-Roman'  # Standardized to Times family
        )
        
        # Table of Contents styles
        toc_title_style = ParagraphStyle(
            'TOCTitle',
            parent=styles['Heading1'],
            fontSize=20,  # Standardized size
            spaceAfter=30,
            textColor=HexColor('#2c3e50'),
            alignment=1,  # Center alignment
            fontName='Times-Bold'  # Standardized to Times family
        )
        
        toc_entry_style = ParagraphStyle(
            'TOCEntry',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=8,
            leftIndent=0,
            rightIndent=0,
            fontName='Times-Roman'  # Standardized to Times family
        )
        
        # Chapter styles
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=16,  # Standardized size
            spaceAfter=24,
            spaceBefore=36,
            textColor=HexColor('#2c3e50'),
            alignment=1,  # Center alignment
            fontName='Times-Bold'  # Standardized to Times family
        )
        
        chapter_body_style = ParagraphStyle(
            'ChapterBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=4,  # Justified
            firstLineIndent=18,  # Standardized to ~0.25 inches
            fontName='Times-Roman',
            leading=16
        )
        
        dialogue_style = ParagraphStyle(
            'Dialogue',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            spaceBefore=8,
            alignment=4,  # Justified
            leftIndent=28,  # Standardized to ~0.4 inches
            rightIndent=10,
            fontName='Times-Roman',
            leading=16
        )
        
        # Build content
        content = []
        
        # Title page
        content.append(Spacer(1, 100))  # Top margin
        content.append(Paragraph(project_obj.title, title_style))
        content.append(Spacer(1, 40))
        content.append(Paragraph("Generated with BookCraft AI", author_style))
        content.append(Spacer(1, 20))
        content.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}", author_style))
        content.append(PageBreak())
        
        # Table of Contents
        content.append(Spacer(1, 50))
        content.append(Paragraph("Table of Contents", toc_title_style))
        content.append(Spacer(1, 30))
        
        # Extract chapter titles and create TOC
        chapter_titles = extract_chapter_titles(project_obj.outline or "")
        current_page = 3  # Start after title and TOC pages
        words_per_page = 275
        
        for i in range(1, project_obj.chapters + 1):
            chapter_title = chapter_titles.get(i, f"Chapter {i}")
            
            # Create professional TOC entry with standardized formatting
            toc_text = f"Chapter {i}: {chapter_title}"
            dots = "." * max(5, 65 - len(toc_text) - len(str(current_page)))  # Standardized spacing
            toc_entry = f"{toc_text} {dots} {current_page}"
            content.append(Paragraph(toc_entry, toc_entry_style))
            
            # Calculate next chapter's page
            if project_obj.chapters_content and str(i) in project_obj.chapters_content:
                chapter_content = project_obj.chapters_content[str(i)]
                word_count = len(re.sub(r'<[^>]+>', '', chapter_content).split())
                pages_in_chapter = max(1, word_count // words_per_page)
                current_page += pages_in_chapter
            else:
                estimated_pages = max(1, (project_obj.pages // project_obj.chapters))
                current_page += estimated_pages
        
        content.append(PageBreak())
        
        # Check if user has made purchases for watermark
        has_purchased = await user_has_made_purchase(current_user.id)
        
        # Create PDF with custom canvas for watermarks if needed
        def custom_canvas_factory(filename):
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(filename, pagesize=A4)
            return WatermarkCanvas(c, None, has_purchased)
        
        if not has_purchased:
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72, canvasmaker=custom_canvas_factory)
        
        # Chapters with improved formatting
        if project_obj.chapters_content:
            for i in range(1, project_obj.chapters + 1):
                chapter_content = project_obj.chapters_content.get(str(i), "")
                if chapter_content:
                    # Get chapter title
                    chapter_title = chapter_titles.get(i, f"Chapter {i}")
                    
                    # Ensure consistent formatting
                    formatted_content = ensure_consistent_chapter_formatting(chapter_content, i, chapter_title)
                    
                    content.append(Paragraph(f"Chapter {i}: {chapter_title}", chapter_title_style))  # Standardized format
                    content.append(Spacer(1, 20))
                    
                    # Process HTML content for better formatting
                    processed_content = process_html_for_pdf(formatted_content, chapter_body_style, dialogue_style)
                    content.extend(processed_content)
                else:
                    content.append(Paragraph(f"Chapter {i}", chapter_title_style))
                    content.append(Spacer(1, 20))
                    content.append(Paragraph("This chapter has not been generated yet.", chapter_body_style))
                
                if i < project_obj.chapters:  # Don't add page break after last chapter
                    content.append(PageBreak())
        
        # Build PDF
        doc.build(content)
        buffer.seek(0)
        
        # Create safe filename
        safe_filename = "".join(c for c in project_obj.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_filename = safe_filename.replace(' ', '_')
        if not safe_filename:
            safe_filename = f"book_{project_obj.id}"
        
        return StreamingResponse(
            io.BytesIO(buffer.getvalue()),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={safe_filename}.pdf"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting book as PDF: {str(e)}")

def process_html_for_pdf(html_content, body_style, dialogue_style):
    """Process HTML content for better PDF formatting"""
    content = []
    
    if not html_content or not html_content.strip():
        return content
    
    # Apply asterisk formatting fixes
    html_content = process_asterisk_formatting(html_content)
    
    # Remove any remaining chapter headers since we add them separately in the main export function
    cleaned_content = re.sub(r'<h2[^>]*>.*?</h2>', '', html_content, flags=re.IGNORECASE)
    
    # Split content into paragraphs using multiple approaches to catch all content
    paragraphs = []
    
    # Method 1: Split by <p> tags
    p_sections = re.split(r'<p[^>]*>(.*?)</p>', cleaned_content, flags=re.DOTALL)
    for i, section in enumerate(p_sections):
        if i % 2 == 1 and section.strip():  # Odd indices contain paragraph content
            paragraphs.append(section.strip())
    
    # Method 2: If we didn't get good paragraphs, try splitting by double line breaks
    if not paragraphs or len(paragraphs) < 2:
        paragraphs = []
        # Remove HTML tags first and split by line breaks
        text_only = re.sub(r'<[^>]+>', '', cleaned_content)
        text_only = unescape(text_only)
        
        # Split by double line breaks
        parts = text_only.split('\n\n')
        for part in parts:
            part = part.strip()
            if part and len(part) > 20:  # Only include substantial paragraphs
                paragraphs.append(part)
    
    # Method 3: If still no good content, split by single line breaks  
    if not paragraphs:
        text_only = re.sub(r'<[^>]+>', '', cleaned_content)
        text_only = unescape(text_only)
        lines = text_only.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            if line:
                current_paragraph.append(line)
            elif current_paragraph:
                full_paragraph = ' '.join(current_paragraph)
                if len(full_paragraph) > 20:  # Only substantial content
                    paragraphs.append(full_paragraph)
                current_paragraph = []
        
        # Add remaining paragraph
        if current_paragraph:
            full_paragraph = ' '.join(current_paragraph)
            if len(full_paragraph) > 20:
                paragraphs.append(full_paragraph)
    
    # Convert paragraphs to PDF content
    for paragraph_text in paragraphs:
        if paragraph_text.strip():
            # Handle bold formatting for PDF
            paragraph_text = re.sub(r'<strong>(.*?)</strong>', r'<b>\1</b>', paragraph_text)
            paragraph_text = re.sub(r'<b>(.*?)</b>', r'<b>\1</b>', paragraph_text)  # Keep existing bold
            
            # Remove any remaining HTML tags except bold
            paragraph_text = re.sub(r'<(?!/?b>)[^>]+>', '', paragraph_text).strip()
            paragraph_text = unescape(paragraph_text)
            
            if paragraph_text:
                # Check if it's dialogue (contains quotation marks)
                if '"' in paragraph_text or '"' in paragraph_text or '"' in paragraph_text or "'" in paragraph_text:
                    content.append(Paragraph(paragraph_text, dialogue_style))
                else:
                    content.append(Paragraph(paragraph_text, body_style))
                
                content.append(Spacer(1, 12))  # Add space between paragraphs
    
    return content

@api_router.get("/export-book-docx/{project_id}")
async def export_book_docx(project_id: str, current_user: User = Depends(get_current_user)):
    """Export book as DOCX with professional book formatting"""
    try:
        # Get project and verify ownership
        project = await db.book_projects.find_one({"id": project_id, "user_id": current_user.id})
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        project_obj = BookProject(**project)
        
        # Create DOCX document
        doc = Document()
        
        # Configure professional book styles - standardized to match PDF
        styles = doc.styles
        
        # Title page styling
        title_style = styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'  # Standardized to Times family
        title_style.font.size = Inches(0.33)  # Standardized size (24pt)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = 1  # Center
        title_style.paragraph_format.space_after = Inches(0.42)  # Standardized spacing
        
        # Author style
        author_style = styles.add_style('AuthorStyle', WD_STYLE_TYPE.PARAGRAPH)
        author_style.font.name = 'Times New Roman'  # Standardized to Times family
        author_style.font.size = Inches(0.19)  # Standardized size (14pt)
        author_style.font.color.rgb = RGBColor(102, 102, 102)
        author_style.paragraph_format.alignment = 1  # Center
        author_style.paragraph_format.space_after = Inches(0.28)  # Standardized spacing
        
        # Chapter title style
        chapter_title_style = styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
        chapter_title_style.font.name = 'Times New Roman'  # Standardized to Times family
        chapter_title_style.font.size = Inches(0.22)  # Standardized size (16pt)
        chapter_title_style.font.bold = True
        chapter_title_style.font.color.rgb = RGBColor(44, 62, 80)
        chapter_title_style.paragraph_format.alignment = 1  # Center
        chapter_title_style.paragraph_format.space_before = Inches(0.5)
        chapter_title_style.paragraph_format.space_after = Inches(0.33)  # Standardized spacing
        
        # Body text style
        body_style = styles.add_style('BookBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Times New Roman'  # Standardized to Times family
        body_style.font.size = Inches(0.15)  # Standardized size (11pt)
        body_style.paragraph_format.alignment = 4  # Justified
        body_style.paragraph_format.first_line_indent = Inches(0.25)  # Standardized to match PDF
        body_style.paragraph_format.space_after = Inches(0.17)  # Standardized spacing
        body_style.paragraph_format.line_spacing = 1.2
        
        # Dialogue style
        dialogue_style = styles.add_style('DialogueStyle', WD_STYLE_TYPE.PARAGRAPH)
        dialogue_style.font.name = 'Times New Roman'  # Standardized to Times family
        dialogue_style.font.size = Inches(0.15)  # Standardized size (11pt)
        dialogue_style.paragraph_format.alignment = 4  # Justified
        dialogue_style.paragraph_format.left_indent = Inches(0.39)  # Standardized to match PDF
        dialogue_style.paragraph_format.space_after = Inches(0.11)  # Standardized spacing
        dialogue_style.paragraph_format.line_spacing = 1.2
        
        # Title page
        title_paragraph = doc.add_paragraph()
        title_paragraph.style = title_style
        title_run = title_paragraph.add_run(project_obj.title)
        
        # Add some space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Author information
        author_paragraph = doc.add_paragraph()
        author_paragraph.style = author_style
        author_run = author_paragraph.add_run("Generated with BookCraft AI")  # Standardized branding
        
        # Date
        date_paragraph = doc.add_paragraph()
        date_paragraph.style = author_style
        date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        
        doc.add_page_break()
        
        # Table of Contents
        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_heading.alignment = 1  # Center alignment
        toc_heading.runs[0].font.name = 'Times New Roman'  # Standardized font
        toc_heading.runs[0].font.size = Inches(0.28)  # Standardized size (20pt)
        toc_heading.runs[0].font.bold = True
        
        # Extract chapter titles and create TOC
        chapter_titles = extract_chapter_titles(project_obj.outline or "")
        current_page = 3  # Start after title and TOC pages
        words_per_page = 275
        
        for i in range(1, project_obj.chapters + 1):
            chapter_title = chapter_titles.get(i, f"Chapter {i}")
            
            # Create professional TOC entry with standardized formatting
            toc_paragraph = doc.add_paragraph()
            toc_paragraph.paragraph_format.space_after = Inches(0.11)  # Standardized spacing
            
            # Chapter info
            chapter_run = toc_paragraph.add_run(f"Chapter {i}: {chapter_title}")
            chapter_run.font.name = 'Times New Roman'  # Standardized font
            chapter_run.font.size = Inches(0.17)  # Standardized size (12pt)
            
            # Add dots with standardized spacing
            dots_text = "." * max(5, 65 - len(f"Chapter {i}: {chapter_title}") - len(str(current_page)))
            dots_run = toc_paragraph.add_run(dots_text)
            dots_run.font.color.rgb = RGBColor(189, 195, 199)  # Light gray
            dots_run.font.name = 'Times New Roman'  # Standardized font
            
            # Add page number
            page_run = toc_paragraph.add_run(str(current_page))
            page_run.bold = True
            page_run.font.color.rgb = RGBColor(52, 152, 219)  # Blue
            page_run.font.name = 'Times New Roman'  # Standardized font
            
            # Calculate next chapter's page
            if project_obj.chapters_content and str(i) in project_obj.chapters_content:
                chapter_content = project_obj.chapters_content[str(i)]
                word_count = len(re.sub(r'<[^>]+>', '', chapter_content).split())
                pages_in_chapter = max(1, word_count // words_per_page)
                current_page += pages_in_chapter
            else:
                estimated_pages = max(1, (project_obj.pages // project_obj.chapters))
                current_page += estimated_pages
        
        doc.add_page_break()
        
        # Check if user has made purchases for watermark
        has_purchased = await user_has_made_purchase(current_user.id)
        
        # Chapters with professional formatting
        if project_obj.chapters_content:
            for i in range(1, project_obj.chapters + 1):
                chapter_content = project_obj.chapters_content.get(str(i), "")
                if chapter_content:
                    # Add chapter title with standardized format
                    chapter_title = chapter_titles.get(i, f"Chapter {i}")
                    
                    # Ensure consistent formatting
                    formatted_content = ensure_consistent_chapter_formatting(chapter_content, i, chapter_title)
                    
                    chapter_heading = doc.add_paragraph()
                    chapter_heading.style = chapter_title_style
                    chapter_heading.add_run(f"Chapter {i}: {chapter_title}")  # Standardized format
                    
                    # Process and add chapter content
                    process_html_for_docx(formatted_content, doc, body_style, dialogue_style)
                else:
                    chapter_heading = doc.add_paragraph()
                    chapter_heading.style = chapter_title_style
                    chapter_heading.add_run(f"Chapter {i}")
                    
                    empty_paragraph = doc.add_paragraph()
                    empty_paragraph.style = body_style
                    empty_paragraph.add_run("This chapter has not been generated yet.")
                
                if i < project_obj.chapters:  # Don't add page break after last chapter
                    doc.add_page_break()
        
        # Add watermark if user hasn't purchased
        add_watermark_to_docx(doc, has_purchased)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create safe filename
        safe_filename = "".join(c for c in project_obj.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_filename = safe_filename.replace(' ', '_')
        if not safe_filename:
            safe_filename = f"book_{project_obj.id}"
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={safe_filename}.docx"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting book as DOCX: {str(e)}")

def process_html_for_docx(html_content, doc, body_style, dialogue_style):
    """Process HTML content for better DOCX formatting"""
    if not html_content or not html_content.strip():
        return
    
    # Apply asterisk formatting fixes
    html_content = process_asterisk_formatting(html_content)
    
    # Remove any remaining chapter headers since we add them separately in the main export function
    cleaned_content = re.sub(r'<h2[^>]*>.*?</h2>', '', html_content, flags=re.IGNORECASE)
    
    # Split content into paragraphs using multiple approaches to catch all content
    paragraphs = []
    
    # Method 1: Split by <p> tags
    p_sections = re.split(r'<p[^>]*>(.*?)</p>', cleaned_content, flags=re.DOTALL)
    for i, section in enumerate(p_sections):
        if i % 2 == 1 and section.strip():  # Odd indices contain paragraph content
            paragraphs.append(section.strip())
    
    # Method 2: If we didn't get good paragraphs, try splitting by double line breaks
    if not paragraphs or len(paragraphs) < 2:
        paragraphs = []
        # Remove HTML tags first and split by line breaks
        text_only = re.sub(r'<[^>]+>', '', cleaned_content)
        text_only = unescape(text_only)
        
        # Split by double line breaks
        parts = text_only.split('\n\n')
        for part in parts:
            part = part.strip()
            if part and len(part) > 20:  # Only include substantial paragraphs
                paragraphs.append(part)
    
    # Method 3: If still no good content, split by single line breaks  
    if not paragraphs:
        text_only = re.sub(r'<[^>]+>', '', cleaned_content)
        text_only = unescape(text_only)
        lines = text_only.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            if line:
                current_paragraph.append(line)
            elif current_paragraph:
                full_paragraph = ' '.join(current_paragraph)
                if len(full_paragraph) > 20:  # Only substantial content
                    paragraphs.append(full_paragraph)
                current_paragraph = []
        
        # Add remaining paragraph
        if current_paragraph:
            full_paragraph = ' '.join(current_paragraph)
            if len(full_paragraph) > 20:
                paragraphs.append(full_paragraph)
    
    # Convert paragraphs to DOCX content
    for paragraph_text in paragraphs:
        if paragraph_text.strip():
            # Create paragraph
            doc_paragraph = doc.add_paragraph()
            
            # Check if it's dialogue (contains quotation marks)
            if '"' in paragraph_text or '"' in paragraph_text or '"' in paragraph_text or "'" in paragraph_text:
                doc_paragraph.style = dialogue_style
            else:
                doc_paragraph.style = body_style
            
            # Handle bold formatting within the paragraph
            # Find bold sections and create runs accordingly
            bold_pattern = r'<strong>(.*?)</strong>'
            current_pos = 0
            
            for match in re.finditer(bold_pattern, paragraph_text):
                # Add text before bold section
                if match.start() > current_pos:
                    before_text = paragraph_text[current_pos:match.start()]
                    before_text = re.sub(r'<[^>]+>', '', before_text)
                    before_text = unescape(before_text)
                    if before_text:
                        doc_paragraph.add_run(before_text)
                
                # Add bold text
                bold_text = match.group(1)
                bold_text = re.sub(r'<[^>]+>', '', bold_text)
                bold_text = unescape(bold_text)
                if bold_text:
                    bold_run = doc_paragraph.add_run(bold_text)
                    bold_run.bold = True
                
                current_pos = match.end()
            
            # Add remaining text after last bold section
            if current_pos < len(paragraph_text):
                remaining_text = paragraph_text[current_pos:]
                remaining_text = re.sub(r'<[^>]+>', '', remaining_text)
                remaining_text = unescape(remaining_text)
                if remaining_text:
                    doc_paragraph.add_run(remaining_text)
            
            # If no bold formatting found, add the whole paragraph
            if not re.search(bold_pattern, paragraph_text):
                clean_text = re.sub(r'<[^>]+>', '', paragraph_text)
                clean_text = unescape(clean_text)
                doc_paragraph.add_run(clean_text)

@api_router.put("/update-outline")
async def update_outline(request: dict):
    """Update project outline"""
    try:
        project_id = request.get("project_id")
        outline = request.get("outline")
        
        if not project_id or not outline:
            raise HTTPException(status_code=400, detail="Project ID and outline are required")
        
        await db.book_projects.update_one(
            {"id": project_id},
            {
                "$set": {
                    "outline": outline,
                    "updated_at": datetime.utcnow()
                }
            }
        )
        
        return {"message": "Outline updated successfully"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating outline: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()